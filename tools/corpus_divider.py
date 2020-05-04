from docx import Document
from nltk.tokenize import RegexpTokenizer
import csv
import os
from glob import glob

tokenizer = RegexpTokenizer('\S+')

def fileToString(filename):
    fullText = []
    if filename[-5:]=='.docx':
        doc = Document(filename)
        for para in doc.paragraphs:
            fullText.append(para.text)
    elif filename[-4:]=='.txt':
        with open(filename) as f:
            fullText.append(f.read())
    string = ' '.join(fullText)
    # print(string)
    return string

def splitCorpus(filename):
    f = fileToString(filename)
    tokens = [token for token in tokenizer.tokenize(f)]
    # print("File has",len(tokens),"words")
    value = 150
    # num = input("If you want "+str(value)+" words per text, enter "+str(len(tokens)//value)+". Please enter a number of files you want to create: ")
    num = len(tokens)//value
    # print("You've entered:",int(num))
    wordsperpart = len(tokens)//int(num)

    matching = [tokens.index(s) for s in tokens if "." in s]
    splitlist = [min(matching, key=lambda x:abs(x-wordsperpart*n))+1 for n in range(1,int(num))]
    res = [tokens[i : j] for i, j in zip([0] +
                                    splitlist, splitlist + [None])]

    # uncomment if you want to write separated texts into separated docx files
    # for n in range(num):
    #     newDoc = Document()
    #     newDoc.add_paragraph(' '.join(res[n]))
    #     newDoc.save("file "+str(n+1)+".docx")
    # print("Created a CSV with",num,"rows.")
    # print("Approximate number of words in 1 row:",wordsperpart)
    return res

# dir - path to the folder where are stored all files you need to put into csv; authorname must coincide with the filename; append = 'a' or 'w' depending if you want to create 1 file or append many texts into 1
def createCSV(dir, authorname,append):

    texts = splitCorpus(dir+authorname+'.txt')
    with open(os.path.abspath(os.path.join(os.getcwd()))+"/tools/corpus.csv",append, newline='') as f:
        thewriter = csv.writer(f)

        # thewriter.writerow(['author','text'])
        for i in range(len(texts)):
            thewriter.writerow([authorname,' '.join(texts[i])])
    return 0

def txtFoldersToString(directoryname):

    dir = glob(os.path.expanduser(directoryname)+"*/") #dir is a list with paths to all the subfolders of given directory
    dir.sort() #sort the paths in alphabetical order
    with open(os.path.basename(os.path.normpath(directoryname))+".csv",'w',newline='') as f:
        thewriter = csv.writer(f)
        thewriter.writerow(['author','text'])
    for folder in dir:
        authorname = os.path.basename(os.path.normpath(folder))
        documents = [f[:-4] for f in os.listdir(folder) if f[-4:] == '.txt'] #creating a list of txt filenames in the current folder
        documents.sort(key = lambda x: int(x.split('-')[0])) #sort the file names in alphabetical order considering 9 < 12
        # print("Documents for folder",os.path.basename(os.path.normpath(folder)),"are",documents)
        contents = []
        for document in documents: #creating a list of texts of one folder
            with open(folder+document+'.txt') as f:
                contents.append(f.read())
        with open(os.path.basename(os.path.normpath(directoryname))+".csv",'a',newline='') as f:
            thewriter = csv.writer(f)
            for i in range(len(contents)):
                thewriter.writerow([authorname,contents[i]])
    # print(*contents, sep='\n\n\n')

    return 0

def txtsToCsv(directoryname):
    dir = os.path.expanduser(directoryname)
    authornames = [f[:-4] for f in os.listdir(dir) if f[-4:] == '.txt'] #creating the list of authornames based on file names without .txt
    authornames.sort()
    print("Folder contains ",len(authornames),"documents/authors.")
    # print("Authornames are:",*authornames,sep="\n")
    with open(os.path.abspath(os.path.join(os.getcwd()))+"/tools/corpus.csv",'w',newline='') as f:
        thewriter = csv.writer(f)
        thewriter.writerow(['author','text'])
    for author in authornames:
        createCSV(dir,author,'a')

    return 0

"~/PycharmProjects/SVM/tools/"
path = os.path.abspath(os.path.join(os.getcwd(), os.pardir))
# txtsToCsv(os.path.expanduser(path+"/corpus/"))